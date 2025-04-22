import streamlit as st
from openai import OpenAI
from docx import Document
from fpdf import FPDF
import requests
import os
from dotenv import load_dotenv
from pathlib import Path

# .envファイルを読み込む
load_dotenv()

# OpenAIクライアント初期化
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


# 環境変数からWebhook URLを取得
teams_webhook_url = os.getenv("TEAMS_WEBHOOK_URL")

# 文字起こし関数（新API仕様）
def transcribe_audio(audio_file):
    response = client.audio.transcriptions.create(
        model="whisper-1",
        file=audio_file,
        response_format="text"
            )
    return response

# 要旨とアクション抽出関数
def summarize_text(text):
    prompt = f"""
    以下の会話内容から「要旨」と「次回のアクションプラン」をそれぞれ抽出してください。

    {text}
    """
    try:
        response = client.Chat.Completion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
        )
        output = response.choices[0].message.content.strip()
        if "アクションプラン：" in output:
            summary, actions = output.split("アクションプラン：")
        else:
            summary, actions = output, "未検出"
        return summary.strip(), actions.strip()
    except Exception as e:
        return "要旨の抽出に失敗しました", f"エラー内容: {str(e)}"

# Word出力関数
def export_to_word(summary, actions):
    doc = Document()
    doc.add_heading("議事録", 0)
    doc.add_heading("要旨", level=1)
    doc.add_paragraph(summary)
    doc.add_heading("アクションプラン", level=1)
    doc.add_paragraph(actions)

    path = "minutes.docx"
    doc.save(path)
    return path

# PDF出力関数
def export_to_pdf(summary, actions):
    pdf = FPDF()
    pdf.add_page()
    # ★ 日本語フォントを読み込む（ipagp.ttf を使う例）
    font_path = "ipagp.ttf"  # IPAexフォント（明朝体またはゴシック体）
    pdf.add_font("IPA", "", font_path, uni=True)
    pdf.set_font("IPA", size=12)

    pdf.multi_cell(0, 10, "【要旨】\n" + summary + "\n\n【アクションプラン】\n" + actions)

    path = "minutes.pdf"
    pdf.output(path)
    return path

# Teams送信関数（簡易）
def send_to_teams(file_path):
    webhook_url = os.getenv("TEAMS_WEBHOOK_URL")
    payload = {
        "text": f"議事録ファイルが作成されました：{file_path}"
    }
    requests.post(webhook_url, json=payload)

# ファイル保存先を一時ディレクトリに変更
import tempfile

def export_to_word(summary, actions):
    doc = Document()
    doc.add_heading("議事録", 0)
    doc.add_heading("要旨", level=1)
    doc.add_paragraph(summary)
    doc.add_heading("アクションプラン", level=1)
    doc.add_paragraph(actions)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return tmp.name
    
# Streamlit UI
st.title("音声データから議事録を作るアプリ")

uploaded_file = st.file_uploader("音声ファイルをアップロードしてください", type=["mp3", "wav", "m4a"])
output_format = st.selectbox("出力形式を選択", ["PDF", "Word"])
send_to = st.multiselect("送信先を選択", ["Slack", "Teams"])

if st.button("議事録を作成"):
    if uploaded_file:
        with st.spinner("音声を文字起こし中..."):
            text = transcribe_audio(uploaded_file)
        
        summary, actions = summarize_text(text)

        if output_format == "PDF":
            file_path = export_to_pdf(summary, actions)
        else:
            file_path = export_to_word(summary, actions)

        st.success("議事録が生成されました！")
        with open(file_path, "rb") as f:
            st.download_button("ダウンロード", f, file_name=file_path)

        if "Teams" in send_to:
            send_to_teams(file_path)
    else:
        st.error("音声ファイルをアップロードしてください。")

# CSSで背景とボタンの色を変更（メインファイル内に直接埋め込む）
st.markdown("""
    <style>
    .stApp {
        background-color: #e3f2fd;
    }
    h1 {
        color: #0d47a1;
        font-family: 'Segoe UI', sans-serif;
        text-align: center;
    }
    .stButton > button {
        background-color: #1976d2;
        color: white;
        font-size: 16px;
        padding: 10px 20px;
        border-radius: 8px;
    }
    </style>
""", unsafe_allow_html=True)

# 右上のアイコンを非常時にする
hide_streamlit_style = """
        <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
        .stDeployButton {display: none;}
        .css-1rs6os.edgvbvh3 {display: none;}
        </style>
        """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)
